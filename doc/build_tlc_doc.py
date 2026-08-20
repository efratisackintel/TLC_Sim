#!/usr/bin/env python3
"""Generate a nicely formatted, table- and diagram-rich Word document for TLC.

Produces TLC_Hierarchy.docx with a title page, auto-updating Table of Contents,
matplotlib diagrams, a throughput chart, compact reference tables and an
appendix of the real firmware data tables (expected-TPT, MCS map, columns).

Run build_tlc_assets.py first to (re)generate the PNG assets.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
OUT = os.path.join(HERE, "TLC_Hierarchy.docx")

BLUE = RGBColor(0x00, 0x71, 0xC5)
DARK = RGBColor(0x1F, 0x2A, 0x44)
GREY = RGBColor(0x55, 0x55, 0x55)
HDR_FILL = "0071C5"
ALT_FILL = "EAF3FB"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Segoe UI"
    if color is not None:
        run.font.color.rgb = color


def add_table(doc, headers, rows, widths=None, size=9, hdr_size=9.5, align_center=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        set_cell_bg(hdr[i], HDR_FILL)
        set_cell_text(hdr[i], head, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                      size=hdr_size, align="center" if align_center else "left")
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if r % 2 == 0:
                set_cell_bg(cells[i], ALT_FILL)
            set_cell_text(cells[i], str(val), size=size,
                          align="center" if (align_center and i > 0) else "left")
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE if level <= 1 else DARK
        run.font.name = "Segoe UI Semibold"
    return p


def body(doc, text, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Segoe UI"
    p.paragraph_format.space_after = Pt(6)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    return p


def image(doc, name, width=6.3, cap=None):
    path = os.path.join(ASSETS, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap:
            caption(doc, cap)
    doc.add_paragraph()


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for typ, txt in [("begin", None), ("instr", 'TOC \\o "1-3" \\h \\z \\u'),
                     ("sep", "Right-click \u2192 Update Field to build the contents."),
                     ("end", None)]:
        if typ == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        elif typ == "sep":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "separate")
            t = OxmlElement("w:t"); t.text = txt; el.append(t)
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), typ)
        run._r.append(el)


def tpt(v):
    """Convert Mbps*10 table value to a Mbps string (blank for 0)."""
    return "" if not v else f"{v/10:.0f}"


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)

    # ---------- Title ----------
    for _ in range(4):
        doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("TLC (Transmit Link Control)")
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = BLUE
    r.font.name = "Segoe UI Semibold"
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Hierarchy, Functionality, Inputs, Outputs & Configuration")
    r.font.size = Pt(15); r.font.color.rgb = DARK
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Firmware Rate Scale Manager \u2014 WCD WiFi")
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY
    for _ in range(9):
        doc.add_paragraph()
    dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dt.add_run("Technical Reference  \u2022  2026")
    r.font.size = Pt(11); r.font.color.rgb = GREY
    doc.add_page_break()

    # ---------- TOC ----------
    h(doc, "Table of Contents", 1)
    add_toc(doc)
    doc.add_page_break()

    # ---------- Overview ----------
    h(doc, "Overview", 1)
    body(doc, "TLC (Rate Scale Manager, \u201cTLC Offload\u201d) is the firmware rate-adaptation "
              "engine. It measures link quality and picks the best PHY/MAC transmission parameters "
              "\u2014 rate, bandwidth, guard interval, spatial streams, aggregation and power \u2014 "
              "to maximize throughput while keeping the link reliable.")
    image(doc, "hierarchy.png", 6.2, "Figure 1. TLC spans four layers; control flows down, statistics flow up.")

    # ---------- 1. Hierarchy ----------
    h(doc, "1. Layered Hierarchy", 1)
    add_table(doc, ["Layer", "Component", "Role"],
              [["Host / Driver", "MmacLinkQuality, DataPathTlcConfig",
                "Sends TLC_MNG_CONFIG_CMD (capabilities); receives TLC_MNG_UPDATE_NTFY"],
               ["Upper MAC", "rateScaleMng, rateScaleAggMng",
                "The 'brain': rate / column / BW / GI search + aggregation management"],
               ["Lower MAC", "rateScale, tlc",
                "Maps station\u2194rsId, applies rate to HW retry table, collects TX statistics"],
               ["PHY / HW", "rate_n_flags decoder",
                "Turns the encoded rate into a PHY vector (modulation, NSS, BW, GI, antenna)"]],
              widths=[1.2, 2.3, 3.3])
    image(doc, "feedback.png", 6.2, "Figure 2. Closed-loop adaptation: every statistics report re-tunes the next TX.")

    # ---------- 2. Units ----------
    h(doc, "2. Units & Responsibilities", 1)
    add_table(doc, ["Layer / File", "Responsibility"],
              [["UMAC \u00b7 rateScaleMng.c", "Core algorithm: state machine, column search, up/down-scale, TPC, expected-TPT tables"],
               ["UMAC \u00b7 rateScaleAggMng.c", "Block-ACK / A-MPDU session mgmt (ADDBA/DELBA, window size, agg factor)"],
               ["UMAC \u00b7 rateScaleMng_static.c", "Static tables (expected TPT, columns, XVT fixed-rate path)"],
               ["LMAC \u00b7 rateScale.c", "station\u2194rsId map, applies rate into HW retry table, agg-allowed decision"],
               ["LMAC \u00b7 tlc.c", "Collects per-station TX statistics, AMSDU BA window, sends TLC_STAT_NTFY"],
               ["PHY", "Consumes rate_n_flags; validated against RS_RATE_ERROR_E"],
               ["Driver", "TLC_MNG_CONFIG_CMD out, TLC_MNG_UPDATE_NTFY in (legacy MmacLinkQuality)"]],
              widths=[2.2, 4.6])

    h(doc, "2.1 Simulation units (USFSTL test framework)", 2)
    add_table(doc, ["Sim file", "Emulates"],
              [["usimLibTlc.c (host)", "Host side: send config, set fixed rate, DHC, read notifs, SNR/path-loss sweeps"],
               ["usimLibLmacTlc.c (lmacSim)", "LMAC\u2192UMAC: fabricate TX statistics, model path-loss\u2192rate\u2192TPT, capture LQ command"],
               ["lsimLibTlc.c", "LMAC LQ config: inject a fixed rate table into LMAC"],
               ["umacSimTlcUTest.c", "End-to-end UMAC tests (per-mode configs, fixed rate, TPC, AMSDU, aggregation)"],
               ["umacSimTlcStaticUTest.c", "Static-table / XVT-mapping tests"]],
              widths=[2.1, 4.7])

    # ---------- 3. Inputs ----------
    h(doc, "3. TLC Inputs \u2014 Origin, Refresh & Usage", 1)
    body(doc, "Every input, with where the value comes from, how often it refreshes, and what it drives.")
    add_table(doc, ["Input", "Origin", "Refresh rate", "Really used \u2014 for what"],
              [["txed / acked", "LMAC tlc.c per BA response", "every BA; pushed at \u226520 frames",
                "Raw material for all metrics (not used directly)"],
               ["Success Ratio (SR)", "UMAC: 32768\u00d7succ/att, EWMA-smoothed", "per stats notif (20 / 2000 frames)",
                "PRIMARY metric: up/down-scale, column compare, TPC, AMSDU"],
               ["PER", "derived = 1 \u2212 SR", "with SR", "Force-decrease boundary (15% / 90% thresholds)"],
               ["Average TPT", "SR \u00d7 expectedTpt / 32768", "with SR", "Compares current vs. candidate column"],
               ["Succ/Fail counters", "UMAC totals", "each update; reset on search", "Trigger a search cycle (4500/400; legacy 480/160)"],
               ["Timestamps", "systemTimeGet()", "each update", "EWMA weighting, stale-data guard, ~10 s idle flush"],
               ["RSSI / SNR", "beacon (MAIN_BEACON_DATA_NTFY)", "per beacon (~100 ms)", "SECONDARY: start-rate pick, TPC gating"],
               ["BA / AMSDU window", "TLC_AMSDU_STAT_S (8 BAs)", "per BA", "AMSDU enable (\u226594%) / disable (\u226450%), size"],
               ["TB statistics", "TLC_TB_STAT_S", "per TB (300 ms / 1 s win)", "UL-OFDMA (trigger-based) rate path"],
               ["PBO / SAR / thermal", "notifications", "event-driven (SAR \u22651 s)", "Cap usable rate/antenna; drive TPC back-off"]],
              widths=[1.25, 1.7, 1.45, 2.4], size=8.5, hdr_size=9)

    # ---------- 4. Outputs ----------
    h(doc, "4. TLC Outputs \u2014 Calculation, Activation & Typical Values", 1)
    add_table(doc, ["Output", "How calculated", "When active", "Typical / notes"],
              [["Primary rate", "Highest-avgTPT rate passing no-decrease SR", "always (first HW attempt)",
                "HE/EHT MCS 9\u201313 in good link; upscale \u2264 every 200 ms"],
               ["Retry table", "Ladder: primary\u2192MCS-1\u2192MCS-2\u2192SISO/legacy", "entry N if all < N failed",
                "~16 entries; last = robust legacy"],
               ["Bandwidth", "column/BW search, clamped to caps/punc", "per frame", "20/40/80/160/320 MHz"],
               ["Guard Interval", "column attribute (RS_MNG_GI_E)", "part of active column", "0.8 \u00b5s hi-SR, 3.2 \u00b5s robust"],
               ["NSS / column", "column-search graph; mimo_delimiter", "MIMO if SR+power allow", "MIMO strong link; SISO/STBC edge"],
               ["Reduced power (TPC)", "2 dB steps; SR 95/98/10% gates", "optimal rate, non-legacy, AMSDU \u22650.5 s", "0\u2013several dB; range \u221210\u2026+24 dBm"],
               ["Antenna mask", "best ant per band", "SISO + control frames", "A / B / A+B (CDD)"],
               ["RTS protection", "rtsBitMsk of protected rates", "unless dbg.rtsDisable", "on higher/aggregated rates"],
               ["Aggregation", "per-mode max; duration 5400 \u00b5s", "once BA session up", "64\u2013256 frames, \u22645.4 ms"],
               ["A-MSDU", "enable \u226594% / disable \u226450% SR", "hi-SR aggregating stations", "3.5\u201311 KB per A-MSDU"],
               ["Rate notif", "mirror of primary rate", "if driver requested notif", "informational to host"]],
              widths=[1.25, 2.0, 1.6, 1.95], size=8.5, hdr_size=9)
    image(doc, "retry.png", 5.6, "Figure 3. The retry table is a descending rate ladder; lower entries are tried only on failure.")

    # ---------- 5. Configuration ----------
    h(doc, "5. Configuration & Tunable Knobs", 1)

    h(doc, "5.1 Station configuration \u2014 TLC_MNG_CONFIG_CMD (from driver)", 2)
    add_table(doc, ["Field", "Controls"],
              [["maxChWidth", "Max bandwidth (20\u2026320 MHz)"],
               ["bestSuppMode", "Highest PHY mode (Legacy/HT/VHT/HE/EHT/UHR)"],
               ["chainsEnabled", "Antenna chains (A/B) \u2192 SISO vs. MIMO"],
               ["sgiChWidthSupport", "SGI allowed per bandwidth (HT/VHT)"],
               ["configFlags", "STBC, LDPC, HE-DCM, EHT-DUP, 2\u00d7LTF block, ELR"],
               ["nonHt", "Bitmap of allowed CCK/OFDM legacy rates"],
               ["mcs[nss][bw]", "Per-NSS / per-BW supported MCS bitmap"],
               ["maxMpduLen", "Drives max A-MSDU size (0 = AMSDU off)"],
               ["maxTxOp", "TXOP cap for all ACs (0 = no limit)"]],
              widths=[2.0, 4.8])

    h(doc, "5.2 Runtime debug hooks \u2014 TLC_MNG_DEBUG_CMD / DHC", 2)
    add_table(doc, ["Type", "Effect"],
              [["FIXED_RATE", "Force a fixed rate_n_flags (0 = auto)"],
               ["PARTIAL_FIXED_RATE", "Fix some fields, scale the rest"],
               ["AGG_DURATION_LIM", "A-MPDU duration (100\u20138000 \u00b5s)"],
               ["AGG_FRAME_CNT_LIM", "A-MPDU frame count (1\u201364)"],
               ["TPC_ENABLED", "Enable/disable TPC (default on)"],
               ["TPC_STATS", "Return per-TPC-step histogram"],
               ["RTS_DISABLE", "Disable RTS protection"],
               ["FAST_START", "Start at max configured PHY rate (WFA)"]],
              widths=[2.3, 4.5])
    body(doc, "Internal-only (\u2265128): MAX_BW, AGG_CONTROL, AGG_F_THRESH, 2ANT_LIMIT, FAST_START, "
              "FIXED_AMSDU, BT_SAR_DICTATE_MODE. Global debug state g_rsMng.dbg holds fastStart, "
              "rtsDisable, amsdu.size/tids, maxBw, durationLimit, sarValue.", size=9.5)

    h(doc, "5.3 TPC configuration \u2014 TPC_CONFIG_S", 2)
    add_table(doc, ["Field", "Default", "Meaning"],
              [["tpcEnabled", "Auto", "0 Auto / 1 Force-off / 2 Force-on"],
               ["tpcSrIncrease", "98%", "SR to step up power back-off"],
               ["tpcSrDecrease", "95%", "SR to step down power back-off"],
               ["tpcSrForceDisable", "10%", "SR to jump back to full power"],
               ["tpcPwrBoStepSize", "2 dB", "Back-off step"],
               ["tpcSearchCycleExtraBo", "2 dB", "Extra back-off during search"]],
              widths=[2.2, 1.1, 3.5])

    h(doc, "5.4 Compile-time algorithm constants (_rateScaleMng.h)", 2)
    add_table(doc, ["Constant", "Value", "Role"],
              [["RS_MNG_PERFECT_SR", "95%", "'perfect' link threshold"],
               ["RS_MNG_SR_NO_DECREASE", "90%", "above this, never downscale"],
               ["RS_MNG_SR_FORCE_DECREASE", "15%", "below this, force downscale"],
               ["RS_STAT_THOLD", "20", "frames per stats notif (normal)"],
               ["RS_MNG_UPSCALE_AGG_FRAME_COUNT", "20", "frames in a test window"],
               ["RS_MNG_SHORT_AGG_FRAME_COUNT", "10", "frames in fast-start window"],
               ["RS_MNG_OPTIMAL_RATE_FRAME_COUNT", "2000", "frames between checks at optimal rate"],
               ["RS_MNG_UPSCALE_MAX_FREQUENCY", "200 ms", "min time between upscale attempts"],
               ["RS_MNG_UPSCALE_SEARCH_CYCLE_MAX_FREQ", "300 ms", "min time between search cycles"],
               ["RS_MNG_AMSDU_SR_ENABLE / DISABLE", "94% / 50%", "AMSDU gating"],
               ["RS_MNG_AGG_DURATION_LIMIT", "5400 \u00b5s", "default A-MPDU duration"],
               ["RS_MNG_EWMA_DENOMINATOR", "256", "SR smoothing resolution"]],
              widths=[3.1, 1.3, 2.4], size=8.5)

    # ---------- 6. State machine ----------
    h(doc, "6. Rate-Scale State Machine", 1)
    image(doc, "state.png", 6.2, "Figure 4. Three states: steady rate, column search, and power (TPC) search.")
    add_table(doc, ["State", "Meaning", "Exit"],
              [["STAY_IN_COLUMN", "Steady rate; STAY / UPSCALE / DOWNSCALE", "\u2192 SEARCH on succ/fail/time"],
               ["SEARCH_CYCLE", "Probe columns/BW; compare averageTpt", "\u2192 STAY or \u2192 TPC_SEARCH"],
               ["TPC_SEARCH", "Reduce power at the optimal rate", "\u2192 STAY when done"]],
              widths=[1.9, 3.5, 1.4])

    # ---------- 7. Reference tables (appendix) ----------
    h(doc, "7. Reference Data Tables (Appendix)", 1)

    h(doc, "7.1 MCS \u2014 Modulation & Coding", 2)
    add_table(doc, ["MCS", "Modulation", "Coding", "MCS", "Modulation", "Coding"],
              [["0", "BPSK", "1/2", "8", "256-QAM", "3/4"],
               ["1", "QPSK", "1/2", "9", "256-QAM", "5/6"],
               ["2", "QPSK", "3/4", "10", "1024-QAM", "3/4"],
               ["3", "16-QAM", "1/2", "11", "1024-QAM", "5/6"],
               ["4", "16-QAM", "3/4", "12", "4096-QAM", "3/4"],
               ["5", "64-QAM", "2/3", "13", "4096-QAM", "5/6"],
               ["6", "64-QAM", "3/4", "14", "DUP DCM", "\u2014"],
               ["7", "64-QAM", "5/6", "15", "DCM", "\u2014"]],
              widths=[0.7, 1.4, 0.9, 0.7, 1.4, 0.9], align_center=True)

    h(doc, "7.2 Expected Throughput \u2014 Legacy (CCK / OFDM), effective Mbps", 2)
    add_table(doc, ["Rate", "Mbps", "Rate", "Mbps", "Rate", "Mbps"],
              [["CCK 1M", tpt(7), "OFDM 6M", tpt(38), "OFDM 24M", tpt(135)],
               ["CCK 2M", tpt(12), "OFDM 9M", tpt(58), "OFDM 36M", tpt(176)],
               ["CCK 5.5M", tpt(33), "OFDM 12M", tpt(76), "OFDM 48M", tpt(208)],
               ["CCK 11M", tpt(54), "OFDM 18M", tpt(108), "OFDM 54M", tpt(221)]],
              widths=[1.3, 0.9, 1.3, 0.9, 1.3, 0.9], align_center=True)
    caption(doc, "Values are effective throughput (raw table stores Mbps\u00d710). Source: expectedTptNonHt.")

    h(doc, "7.3 Expected Throughput \u2014 HT/VHT, 80 MHz (effective Mbps)", 2)
    htvht80 = {
        "SISO NGI": [237, 478, 718, 959, 1427, 1849, 2043, 2232, 2589, 2813],
        "SISO SGI": [263, 531, 797, 1065, 1577, 2022, 2231, 2434, 2814, 3052],
        "MIMO NGI": [475, 953, 1420, 1843, 2584, 3229, 3522, 3799, 4303, 4603],
        "MIMO SGI": [528, 1058, 1569, 2016, 2808, 3490, 3798, 4087, 4610, 4919],
    }
    rows = [[k] + [tpt(v) for v in vals] for k, vals in htvht80.items()]
    add_table(doc, ["Config"] + [f"M{i}" for i in range(10)], rows,
              widths=[1.1] + [0.57] * 10, size=8, hdr_size=8, align_center=True)
    caption(doc, "MCS 0\u20139, 80 MHz. Source: expectedTptHtVht[CHANNEL_WIDTH80].")

    h(doc, "7.4 Expected Throughput \u2014 HE, 80 MHz (effective Mbps)", 2)
    he80 = {
        "SISO 3.2\u00b5s": [260, 520, 780, 1041, 1561, 2082, 2342, 2603, 3123, 3470, 3904, 4338],
        "SISO 1.6\u00b5s": [289, 578, 867, 1156, 1735, 2313, 2603, 2892, 3470, 3856, 4338, 4820],
        "SISO 0.8\u00b5s": [306, 612, 918, 1225, 1837, 2450, 2756, 3062, 3675, 4083, 4593, 5104],
        "MIMO 3.2\u00b5s": [520, 1041, 1561, 2082, 3123, 4165, 4685, 5206, 6247, 6941, 7809, 8677],
        "MIMO 1.6\u00b5s": [578, 1156, 1735, 2313, 3470, 4627, 5206, 5784, 6941, 7712, 8677, 9641],
        "MIMO 0.8\u00b5s": [612, 1225, 1837, 2450, 3675, 4900, 5512, 6125, 7350, 8166, 9187, 10208],
    }
    rows = [[k] + [tpt(v) for v in vals] for k, vals in he80.items()]
    add_table(doc, ["Config"] + [f"M{i}" for i in range(12)], rows,
              widths=[1.0] + [0.48] * 12, size=7.5, hdr_size=7.5, align_center=True)
    caption(doc, "MCS 0\u201311, 80 MHz. Source: expectedTptHe[CHANNEL_WIDTH80].")
    image(doc, "tpt_chart.png", 6.2, "Figure 5. HE 80 MHz, GI 0.8 \u00b5s \u2014 MIMO roughly doubles SISO throughput.")

    h(doc, "7.5 Rate-Scale Columns (RS_MNG_COLUMN_DESC_E)", 2)
    add_table(doc, ["Group", "Columns"],
              [["Legacy (CCK/OFDM)", "NON_HT_ANT_A, NON_HT_ANT_B"],
               ["HT/VHT SISO", "SISO_STBC, SISO_ANT_A, SISO_ANT_B (+ _SGI variants)"],
               ["HT/VHT MIMO", "MIMO, MIMO_SGI"],
               ["HE SISO", "HE_{3.2/1.6/0.8}_SISO_{STBC/ANT_A/ANT_B}"],
               ["HE MIMO", "HE_{3.2/1.6/0.8}_MIMO"]],
              widths=[1.9, 4.9])
    body(doc, "A column = a fixed combination of mode, NSS, GI and antenna. The search cycle walks a "
              "graph of related columns (nextCols) and compares averageTpt to find the best one.", size=9.5)

    # ---------- 8. File map ----------
    h(doc, "8. Quick File Map", 1)
    add_table(doc, ["Layer", "Path"],
              [["UMAC algorithm", "fw/src/umac/main/dataPath/rateScaleMng/"],
               ["LMAC apply", "fw/src/lmac/mcm/core/dataPath/rateScale/"],
               ["LMAC statistics", "fw/src/lmac/mcm/core/dataPath/tlc/"],
               ["API config/notif", "fw/inc/commands/apiGroupDatapath.h"],
               ["API LQ/retry table", "fw/src/shared/inc/apiLmacUmac.h, stationDefines.h"],
               ["Host sim", "softwareTesting/.../umac/libs/host/usimLibTlc.c"],
               ["LMAC sim", "softwareTesting/.../umac/libs/lmacSim/usimLibLmacTlc.c"],
               ["UMAC tests", "softwareTesting/projects/usim/umacSimTlc/umacSimTlcUTest.c"],
               ["Driver config", "wifi_drv-dev/.../DataPath/DataPathTlcConfig/DataPathTlcConfig.c"],
               ["Driver legacy TLC", "wifi_drv-dev/.../MMAC/MmacLinkQuality/"]],
              widths=[1.8, 5.0])

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
